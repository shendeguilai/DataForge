#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import copy
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.table import Table
except ImportError:
    print('缺少依赖 python-docx。请运行：python3 -m pip install python-docx')
    sys.exit(2)

SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = SCRIPT_DIR / 'CSP_初赛模板.docx'


class ConvertError(RuntimeError):
    pass


def info(msg):
    print(f'[CSP] {msg}')


def normalize_math_tex(tex: str) -> str:
    tex = re.sub(r'\\tt\s*\{([^{}]*)\}', r'\\mathtt{\1}', tex)
    tex = re.sub(r'\\tt\s+([A-Za-z0-9_]+)', r'\\mathtt{\1}', tex)
    return tex.strip()


def collect_math(md: str):
    formulas, seen = [], set()
    tmp = re.sub(r'```.*?```', '', md, flags=re.S)
    for m in re.finditer(r'(?<!\\)\$([^$\n]+?)(?<!\\)\$', tmp):
        tex = normalize_math_tex(m.group(1))
        # 题库常把代码变量写成 $cnt\_check$，这种不送进公式转换
        if re.fullmatch(r'[A-Za-z][A-Za-z0-9]*(?:\\_[A-Za-z0-9]+)+', tex):
            continue
        if tex not in seen:
            seen.add(tex)
            formulas.append(tex)
    return formulas


def build_math_bank(formulas):
    if not formulas:
        return {}
    if shutil.which('pandoc') is None:
        raise ConvertError(
            '未找到 pandoc。为了保证上下标/求和等公式正确，请先安装 Pandoc。\n'
            'macOS：brew install pandoc\nWindows：winget install --id JohnMacFarlane.Pandoc'
        )
    with tempfile.TemporaryDirectory(prefix='csp_math_') as td:
        td = Path(td)
        md_path, docx_path = td/'math.md', td/'math.docx'
        lines = []
        for i, tex in enumerate(formulas):
            lines += [f'F{i:04d} ${tex}$', '']
        md_path.write_text('\n'.join(lines), encoding='utf-8')
        cp = subprocess.run(
            ['pandoc', str(md_path), '-f', 'markdown+tex_math_dollars', '-o', str(docx_path)],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        if cp.returncode != 0:
            raise ConvertError('Pandoc 公式转换失败：\n' + cp.stderr)
        bad = [x for x in cp.stderr.splitlines() if 'Could not convert TeX math' in x]
        if bad:
            raise ConvertError('存在 Pandoc 无法识别的公式：\n' + '\n'.join(bad))
        d = Document(docx_path)
        bank = {}
        for i, p in enumerate(d.paragraphs[:len(formulas)]):
            maths = p._p.xpath('.//m:oMath')
            if not maths:
                raise ConvertError(f'公式未成功转换：${formulas[i]}$')
            bank[formulas[i]] = copy.deepcopy(maths[0])
        if len(bank) != len(formulas):
            raise ConvertError('公式转换数量异常。')
        return bank


def split_sections(md: str):
    parts = re.split(r'(?m)^## 第\s*(\d+)\s*题\s*$', md)
    out = {}
    for i in range(1, len(parts), 2):
        out[int(parts[i])] = parts[i+1]
    return out


def clear_paragraph(p):
    for child in list(p._p):
        if child.tag != qn('w:pPr'):
            p._p.remove(child)


def set_run_font(run, name='Times New Roman', east='Noto Serif CJK SC', size=None):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for k, v in [('ascii', name), ('hAnsi', name), ('eastAsia', east)]:
        rFonts.set(qn('w:' + k), v)
    if size is not None:
        run.font.size = Pt(size)


def add_text_run(p, text, code=False):
    if not text:
        return
    r = p.add_run(text)
    set_run_font(r, 'Consolas' if code else 'Times New Roman',
                 'Noto Sans Mono CJK SC' if code else 'Noto Serif CJK SC')


def add_math(p, tex, bank):
    tex = normalize_math_tex(tex)
    if tex not in bank:
        raise ConvertError(f'公式库缺少：${tex}$')
    p._p.append(copy.deepcopy(bank[tex]))


def fill_inline(p, text, bank):
    clear_paragraph(p)
    text = str(text).replace('\\(', '$').replace('\\)', '$')
    pos = 0
    token_re = re.compile(r'(\$[^$\n]*\$|`[^`\n]*`)')
    for m in token_re.finditer(text):
        if m.start() > pos:
            add_text_run(p, text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('$'):
            raw = tok[1:-1].strip()
            if re.fullmatch(r'[A-Za-z][A-Za-z0-9]*(?:\\_[A-Za-z0-9]+)+', raw):
                add_text_run(p, raw.replace('\\_', '_'), code=True)
            else:
                add_math(p, raw, bank)
        else:
            add_text_run(p, tok[1:-1], code=True)
        pos = m.end()
    if pos < len(text):
        add_text_run(p, text[pos:])


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def add_p(doc, style, text='', bank=None, keep_next=False):
    p = doc.add_paragraph(style=style)
    if bank is None:
        p.add_run(text)
    else:
        fill_inline(p, text, bank)
    if keep_next:
        p.paragraph_format.keep_with_next = True
    return p


def clone_table_to_doc(doc, ref_table):
    tbl = copy.deepcopy(ref_table._tbl)
    doc._body._element.insert(-1, tbl)
    return Table(tbl, doc._body)


def flatten_cells(table):
    out, seen = [], set()
    for row in table.rows:
        for c in row.cells:
            key = id(c._tc)
            if key not in seen:
                seen.add(key); out.append(c)
    return out


def set_cell(cell, text, bank):
    p = cell.paragraphs[0]
    try: p.style = 'CSPJ 选项'
    except Exception: pass
    fill_inline(p, text, bank)
    for x in cell.paragraphs[1:]:
        x._element.getparent().remove(x._element)


def visible_len(s):
    s = re.sub(r'`([^`]*)`', r'\1', s)
    s = re.sub(r'\$([^$]*)\$', r'\1', s)
    # 中文按 2，ASCII 按 1；保留给结构检查和旧逻辑兼容。
    return sum(2 if ord(ch) > 127 else 1 for ch in s)


def _plain_width_units(text: str, monospace=False) -> float:
    """估算一段文本的视觉宽度（单位约等于 1 个中文字符宽）。

    这里不是精确排版引擎，只用于在 1x4 / 2x2 / 4x1 三种固定 Word
    选项表之间做更保守的选择。代码使用等宽字体，因此明显比普通英文
    更容易在四栏中换行，必须单独加权。
    """
    width = 0.0
    for ch in text:
        if ch.isspace():
            width += 0.28
        elif ord(ch) > 127:
            width += 1.0
        elif ch.isalnum() or ch == '_':
            width += 0.64 if monospace else 0.54
        else:
            width += 0.55 if monospace else 0.42
    return width


def option_width_units(s: str) -> float:
    """按最终渲染字体粗略估算一个选项的宽度。

    - `code` 片段按 Consolas/等宽字体估算；
    - $math$ 去掉 TeX 控制字符后按普通西文估算；
    - 中文按一个全角字符宽度计算。
    """
    width = 0.0
    pos = 0
    token_re = re.compile(r'(`[^`\n]*`|\$[^$\n]*\$)')
    for m in token_re.finditer(s):
        width += _plain_width_units(s[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('`'):
            width += _plain_width_units(tok[1:-1], monospace=True)
        else:
            tex = tok[1:-1]
            # TeX 命令本身不占那么多可见宽度，例如 \sqrt、\log。
            tex = re.sub(r'\\[A-Za-z]+', 'X', tex)
            tex = tex.replace('{', '').replace('}', '')
            width += _plain_width_units(tex)
        pos = m.end()
    width += _plain_width_units(s[pos:])
    return width


def choose_option_shape(opts):
    """根据最终 Word 单元格的可用宽度选择选项布局。

    目标不是尽量塞满一行，而是尽量避免 Word 自己把一个选项从中间断行。
    模板正文约 12pt，四栏单元格的有效宽度很有限；尤其反引号代码会使用
    Consolas 等宽字体，视觉宽度明显大于普通英文，因此需要比旧版更保守。

    返回：
      - 1x4：四个选项同一行；
      - 2x2：两列两行；
      - 4x1：每个选项独占一行。
    """
    if not opts:
        return '4x1'

    widths = [option_width_units(x) for x in opts]
    mx = max(widths)
    total = sum(widths)

    # 两个选项通常一行两列即可；极长文本仍改为单列，避免跨行后难看。
    if len(opts) <= 2:
        return '2x2' if mx <= 16.2 else '4x1'

    # 四栏：每格约 1/4 版心。经验阈值 7.2 能覆盖纯数字、短公式、短文字，
    # 但会把 n % (i-1) == 0 这类等宽代码主动降为 2x2。
    if mx <= 7.2 and total <= 24.5:
        return '1x4'

    # 两栏：每格约 1/2 版心。超过约 16 个中文字符等效宽度时，Word 很容易
    # 在 CJK/代码混排中换行，因此直接改为单列。
    if mx <= 16.2 and total <= 55.0:
        return '2x2'

    return '4x1'

def _trim_unused_option_rows(table, shape, option_count):
    """删除选项表里完全没有使用的尾部行。

    旧版为了复用固定模板，会把两个判断选项 A/B 填进 2x2 表格的第一行，
    但第二行仍保留为空，于是 Word 中会出现一整行多余空白。

    这里保留模板本身的列宽、边距、字体等样式，只删除真正用不到的尾部行：
      - 2x2 + 2 个选项 -> 只保留第 1 行（实际 1x2）
      - 4x1 + 2/3 个选项 -> 只保留对应的 2/3 行
    """
    if shape == '2x2':
        needed_rows = max(1, (option_count + 1) // 2)
    elif shape == '4x1':
        needed_rows = max(1, option_count)
    else:
        needed_rows = 1

    while len(table.rows) > needed_rows:
        row = table.rows[-1]
        row._tr.getparent().remove(row._tr)


def add_option_table(doc, refs, opts, bank):
    shape = choose_option_shape(opts)
    ref = refs[shape]
    t = clone_table_to_doc(doc, ref)
    _trim_unused_option_rows(t, shape, len(opts))
    cells = flatten_cells(t)
    if len(cells) < len(opts):
        raise ConvertError(f'选项表容量不足：{shape}')
    for i, x in enumerate(opts):
        set_cell(cells[i], x, bank)
    for c in cells[len(opts):]:
        set_cell(c, '', bank)
    return t


def normalize_code_lines(code):
    """去掉题库 Markdown 代码块中已经存在的行号，同时保留代码缩进。

    部分年份（尤其 2022 CSP-J）导出的代码块本身已经带行号，例如：
        01 #include <iostream>
        07     unsigned short x, y;

    旧版直接用正则把“行号后的全部空白”都吃掉，导致第二行原有的 4 个
    缩进空格也被删除。这里把源行号视为一个固定宽度的“行号列”：
    行号列宽 = 本代码块最大行号位数 + 1 个分隔空格。

    例如最大行号是 49，则：
        '1  #include'  -> '#include'       （删除 2 个分隔空格）
        '10 int h...'  -> 'int h...'       （删除 1 个分隔空格）
        '14     if...' -> '    if...'      （只删 1 个分隔空格，保留 4 空格缩进）
    """
    raw = code.rstrip('\n').splitlines() or ['']
    nonempty = [x for x in raw if x.strip()]
    if not nonempty:
        return raw

    parsed = []
    numbered = 0
    max_digits = 0
    for x in nonempty:
        m = re.match(r'^\s*(\d{1,3})([ \t]*)(.*)$', x)
        if m:
            numbered += 1
            max_digits = max(max_digits, len(m.group(1)))
            parsed.append(m)

    # 题库行号块通常几乎每行都有编号；80% 阈值避免普通代码被误判。
    if numbered / len(nonempty) < 0.80 or max_digits == 0:
        return raw

    prefix_width = max_digits + 1
    out = []
    for x in raw:
        m = re.match(r'^\s*(\d{1,3})([ \t]*)(.*)$', x)
        if not m:
            out.append(x)
            continue

        num, ws, rest = m.group(1), m.group(2), m.group(3)
        if not rest:
            # 只有行号的空行，例如“02”或“16”。
            out.append('')
            continue

        # 只移除行号列为对齐所占的空白，额外空白就是代码自身缩进。
        need_remove = max(0, prefix_width - len(num))
        if '\t' not in ws:
            kept_ws = ws[min(need_remove, len(ws)):]
        else:
            # 极少数源文件可能使用 Tab 分隔；逐字符删除“分隔部分”，
            # 不展开 Tab，避免改变原代码的相对缩进。
            kept_ws = ws
            removed = 0
            while kept_ws and removed < need_remove:
                kept_ws = kept_ws[1:]
                removed += 1
        out.append(kept_ws + rest)
    return out


def set_code_table(table, code):
    lines = normalize_code_lines(code)

    # 代码超过模板预留行数时，不能直接 table.add_row()：
    # python-docx 新增的空行不会完整继承模板中的单元格宽度、段落缩进、
    # 边框等格式，导致代码中途（例如第 29 行）突然变样。
    # 这里直接复制模板最后一行，保证任意长度代码都保持同一格式。
    while len(table.rows) < len(lines):
        table._tbl.append(copy.deepcopy(table.rows[-1]._tr))
    while len(table.rows) > len(lines):
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)

    for i, line in enumerate(lines, 1):
        row = table.rows[i-1]
        for ci, txt in enumerate((f'{i:02d}', line)):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            try: p.style = 'CSPJ 代码'
            except Exception: pass
            clear_paragraph(p)
            # 无论来自模板原有行还是动态复制行，都显式统一段落格式。
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt)
            set_run_font(r, 'Consolas', 'Noto Sans Mono CJK SC', 9.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_code_table(doc, refs, code):
    t = clone_table_to_doc(doc, refs['code'])
    set_code_table(t, code)
    return t


def lines_options(text):
    opts = []
    for line in text.splitlines():
        m = re.match(r'^\s*-?\s*([A-D])\.\s*(.+?)\s*$', line)
        if m:
            opts.append(m.group(1)+'. '+m.group(2))
    return opts


def first_abcd(opts):
    first, expected = [], 'A'
    for opt in opts:
        if opt.startswith(expected+'.'):
            first.append(opt)
            expected = chr(ord(expected)+1)
            if expected == 'E': break
        elif opt.startswith('A.'):
            first, expected = [opt], 'B'
    return first


def parse_q1_15(sec, n):
    """解析前 15 道单项选择题。

    洛谷不同年份的 MD 并不完全一致：
    - 常见格式：`第 2 题  题干...`
    - 2023 第 1 题：先写 `### 一、单项选择题...`，下一段直接是题干，
      没有 `第 1 题` 前缀；
    - 有些单选题题干中还夹有 cpp 代码块。

    因此这里不再依赖“第 n 题”这几个字，而是以 A 选项为边界，
    把前面的内容按正文/代码块顺序解析。
    """
    before = sec.split('本题共', 1)[0]
    am = re.search(r'(?m)^\s*-\s*A\.\s*', before)
    if not am:
        raise ConvertError(f'第 {n} 题没有找到 A 选项，无法确定题干边界。')

    stem = before[:am.start()]
    # 删除该 section 内的 Markdown 小标题，例如 2023 第 1 题中的
    # `### 一、单项选择题...`，以及其它年份可能出现的辅助标题。
    stem = re.sub(r'(?m)^\s*#{3,6}\s+.*?\s*$', '', stem)

    blocks = []
    pos = 0
    for cm in re.finditer(r'```cpp\s*\n(.*?)\n```', stem, re.S):
        txt = stem[pos:cm.start()]
        for para in re.split(r'\n\s*\n', txt):
            para = re.sub(r'\s*\n\s*', ' ', para.strip())
            if para:
                blocks.append(('text', para))
        blocks.append(('code', cm.group(1)))
        pos = cm.end()
    tail = stem[pos:]
    for para in re.split(r'\n\s*\n', tail):
        para = re.sub(r'\s*\n\s*', ' ', para.strip())
        if para:
            blocks.append(('text', para))

    # 常规年份题干开头会重复“第 n 题”，统一去掉；没有该前缀也没关系。
    for i, (kind, val) in enumerate(blocks):
        if kind == 'text':
            val = re.sub(rf'^第\s*{n}\s*题\s*', '', val).strip()
            blocks[i] = (kind, val)
            break

    blocks = [(k, v) for k, v in blocks if v.strip()]
    if not blocks:
        raise ConvertError(f'第 {n} 题没有找到题干。')

    opts = first_abcd(lines_options(before))
    if len(opts) != 4:
        raise ConvertError(f'第 {n} 题应有 A-D 四个选项，实际 {len(opts)}。')
    return blocks, opts


def strip_dup_tail(text):
    return re.split(r'(?m)^\s*-\s*1\.\s*$', text, maxsplit=1)[0]


def duplicate_groups(text):
    """读取洛谷导出尾部的 - 1. / A. ... / B. ... 结构。"""
    m = re.search(r'(?m)^\s*-\s*1\.\s*$', text)
    if not m:
        return []
    tail = text[m.start():]
    starts = list(re.finditer(r'(?m)^\s*-\s*(\d+)\.\s*$', tail))
    groups = []
    for i, sm in enumerate(starts):
        st = sm.end(); ed = starts[i+1].start() if i+1 < len(starts) else len(tail)
        block = tail[st:ed]
        opts = lines_options(block)
        groups.append((int(sm.group(1)), first_abcd(opts)))
    return groups


def bullet_questions(text):
    out = []
    for line in text.splitlines():
        m = re.match(r'^\s*-\s+(.+?)\s*$', line)
        if not m: continue
        s = m.group(1).strip()
        if re.match(r'^[A-D]\.', s): continue
        if re.match(r'^\d+\.\s*$', s): continue
        out.append(s)
    return out


def parse_reading(sec, qno):
    cm = re.search(r'```cpp\s*\n(.*?)\n```', sec, re.S)
    if not cm:
        raise ConvertError(f'第 {qno} 题没有 cpp 代码块。')
    code = cm.group(1)
    predup = strip_dup_tail(sec)

    # 不再要求“判断题 / 单选题”必须写成 #### 标题。
    # 支持：
    #   #### 判断题
    #   #### 单选题
    # 以及 2022 洛谷导出的普通文本行：
    #   判断题
    #   单选题
    # 只匹配“独立一整行”，不会误伤“完成下面的判断题和单选题”这类说明文字。
    judge_heads = list(re.finditer(
        r'(?mi)^\s*(?:#{1,6}\s*)?判断题\s*$', predup
    ))
    choice_heads = list(re.finditer(
        r'(?mi)^\s*(?:#{1,6}\s*)?(?:单选题|选择题)\s*$', predup
    ))
    if not judge_heads or not choice_heads:
        raise ConvertError(f'第 {qno} 题没有识别到独立的“判断题”和“单选题/选择题”分隔行。')

    # 取代码块之后出现的第一组“判断题 -> 单选题/选择题”。
    jh = next((x for x in judge_heads if x.start() > cm.end()), None)
    if jh is None:
        jh = judge_heads[0]
    sh = next((x for x in choice_heads if x.start() > jh.end()), None)
    if sh is None:
        raise ConvertError(f'第 {qno} 题识别到“判断题”，但其后没有“单选题/选择题”。')

    judge_text = predup[jh.end():sh.start()]
    select_text = predup[sh.end():]
    judges = bullet_questions(judge_text)
    selects = bullet_questions(select_text)
    if not judges:
        raise ConvertError(f'第 {qno} 题没有解析到判断题。')
    if not selects:
        raise ConvertError(f'第 {qno} 题没有解析到单选题。')

    dups = duplicate_groups(sec)
    # 选择题选项优先从重复答案区读取，兼容“题干和选项分离”的洛谷 MD。
    choice_opts = []
    for idx in range(len(judges), len(judges)+len(selects)):
        if idx < len(dups) and len(dups[idx][1]) == 4:
            choice_opts.append(dups[idx][1])
    if len(choice_opts) != len(selects):
        # 回退：尝试直接从单选题区按“题干 + A-D”解析
        items = []
        for line in select_text.splitlines():
            mm = re.match(r'^\s*-\s+(.+?)\s*$', line)
            if mm: items.append(mm.group(1).strip())
        groups, i = [], 0
        while i < len(items):
            if not re.match(r'^[A-D]\.', items[i]):
                q = items[i]; opts=[]; j=i+1
                while j<len(items) and re.match(r'^[A-D]\.', items[j]) and len(opts)<4:
                    opts.append(items[j]); j+=1
                if len(opts)==4: groups.append((q, opts)); i=j; continue
            i += 1
        if len(groups)==len(selects):
            selects = [x[0] for x in groups]
            choice_opts = [x[1] for x in groups]
        else:
            raise ConvertError(f'第 {qno} 题有 {len(selects)} 道单选题，但只解析到 {len(choice_opts)} 组选项。')

    # 代码块结束后、判断题标题之前的文字作为阅读程序说明。
    note = predup[cm.end():jh.start()].strip()
    # 去掉可能夹在其中的小标题（如 #### (1)），但保留普通说明文字。
    note = re.sub(r'(?m)^\s*#{1,6}\s+.*?\s*$', '', note).strip()
    return code, judges, list(zip(selects, choice_opts)), note


def parse_fill(sec, qno):
    cm = re.search(r'```cpp\s*\n(.*?)\n```', sec, re.S)
    if not cm:
        raise ConvertError(f'第 {qno} 题没有 cpp 代码块。')
    code = cm.group(1)
    pre = sec[:cm.start()]
    post = strip_dup_tail(sec[cm.end():])
    items=[]
    for line in post.splitlines():
        m=re.match(r'^\s*-\s+(.+?)\s*$',line)
        if m: items.append(m.group(1).strip())
    groups=[]; i=0
    while i<len(items):
        if not re.match(r'^[A-D]\.',items[i]):
            q=items[i]; opts=[]; j=i+1
            while j<len(items) and re.match(r'^[A-D]\.',items[j]) and len(opts)<4:
                opts.append(items[j]); j+=1
            if len(opts)==4:
                groups.append((q,opts)); i=j; continue
        i+=1

    dups=duplicate_groups(sec)
    if not groups:
        # 常见洛谷格式：题干列表和选项列表分离。
        qs=[x for x in bullet_questions(post) if '处应填' in x]
        if len(qs) >= 5:
            groups=[(q,dups[i][1]) for i,q in enumerate(qs[:5]) if i<len(dups) and len(dups[i][1])==4]
        # 2022 CSP-J 格式只有一行“①~⑤处应填（ ）”，后面直接给 5 组选项。
        # 此时根据原文中的 ①~⑤ 语义拆成 5 个独立题目用于 Word 排版。
        elif len(dups) >= 5 and re.search(r'[①1]\s*[~～-]\s*[⑤5]\s*处应填', post):
            marks='①②③④⑤'
            groups=[]
            for i,mark in enumerate(marks):
                if i < len(dups) and len(dups[i][1]) == 4:
                    groups.append((f'{mark} 处应填（ ）', dups[i][1]))

    if len(groups) != 5:
        raise ConvertError(f'第 {qno} 题应有 5 个填空选择题，实际解析到 {len(groups)}。')
    return pre, code, groups


def parse_meta(md):
    h1=re.search(r'(?m)^#\s+(.+?)\s*$',md)
    title=h1.group(1) if h1 else ''
    ym=re.search(r'20\d{2}',title); year=ym.group(0) if ym else '20XX'
    if '提高级' in title or 'CSP-S' in title.upper(): return year,'S','提高级'
    return year,'J','入门级'


def extract_heading(sec, prefix, default):
    m=re.search(rf'(?m)^####\s*({re.escape(prefix)}[^\n]*)$',sec)
    return m.group(1).strip() if m else default


def clean_prose(pre):
    """把完善程序代码前的 Markdown 说明整理成正文。

    - 大题标题“### 三、完善程序...”不重复输出；
    - 纯编号小标题如“#### 第 1 题”忽略；
    - 2022 的“#### （1）（枚举因数）从小到大...”中，真正的题目说明要保留，
      只去掉 Markdown # 和开头的（1）编号，避免与转换器生成的（1）重复。
    """
    lines=[]
    for ln in pre.splitlines():
        raw=ln.strip()
        hm=re.match(r'^#{1,6}\s*(.*?)\s*$', raw)
        if hm:
            h=hm.group(1).strip()
            if re.match(r'^三[、.]\s*完善程序', h):
                continue
            if re.match(r'^第\s*\d+\s*题\s*$', h):
                continue
            # 去掉本块编号，但保留其后的题目名/说明。
            h=re.sub(r'^[（(]\s*\d+\s*[）)]\s*', '', h).strip()
            if h:
                lines.append(h)
            continue
        if not raw:
            lines.append('')
            continue
        if re.match(r'^\s*-\s+',ln):
            lines.append('• '+re.sub(r'^\s*-\s+','',ln).strip())
        else:
            lines.append(raw)
    text='\n'.join(lines).strip()
    blocks=[]
    for b in re.split(r'\n\s*\n',text):
        b=b.strip()
        if not b: continue
        if b.startswith('• '):
            for x in b.splitlines():
                if x.strip(): blocks.append(x.strip())
        else:
            blocks.append(re.sub(r'\s*\n\s*',' ',b))
    return blocks


def setup_refs(template_doc):
    # 自带模板中这些表分别代表 1x4、2x2、4x1 和代码框的标准样式
    return {'1x4': template_doc.tables[0], '2x2': template_doc.tables[1],
            '4x1': template_doc.tables[34], 'code': template_doc.tables[15]}


def convert(src, out, template, numbering='auto'):
    md=src.read_text(encoding='utf-8-sig')
    secs=split_sections(md)
    missing=[i for i in range(1,21) if i not in secs]
    if missing: raise ConvertError('MD 缺少一级题目段：'+','.join(map(str,missing)))

    formulas=collect_math(md)
    info(f'发现 {len(formulas)} 个数学公式，正在转换为 Word 公式……')
    bank=build_math_bank(formulas)

    template_doc=Document(template)
    if len(template_doc.tables)<36:
        raise ConvertError('模板不符合要求。')
    refs=setup_refs(template_doc)
    doc=Document(template)
    clear_document_body(doc)
    year,group,level=parse_meta(md)

    # 自动模式：J 组旧式试卷使用全局题号；S 组 2025 样式使用块内 (1)(2)...
    if numbering=='auto':
        global_numbering = (group=='J')
    else:
        global_numbering = (numbering=='global')

    add_p(doc,'CSPJ 标题',f'{year} CCF 非专业级别软件能力认证第一轮',bank)
    add_p(doc,'CSPJ 副标题',f'（CSP-{group}1）{level} C++ 语言试题',bank)
    add_p(doc,'CSPJ 考试信息','考试时间：120 分钟    试卷满分：100 分',bank)
    add_p(doc,'Normal','',bank)
    add_p(doc,'CSPJ 注意标题','考生注意事项：',bank,True)
    add_p(doc,'CSPJ 注意事项','• 试题分为单项选择题、阅读程序题和完善程序题三部分。',bank)
    add_p(doc,'CSPJ 注意事项','• 所有程序均使用 C++ 编写；除特殊说明外，程序输入不超过数组或字符串定义的范围。',bank)
    add_p(doc,'CSPJ 注意事项','• 单项选择题和完善程序题每题有且仅有一个正确选项。',bank)
    if group=='J':
        add_p(doc,'CSPJ 注意事项','• 阅读程序中的判断题，正确填“√”，错误填“×”。',bank)
    else:
        add_p(doc,'CSPJ 注意事项','• 程序阅读题中的判断题，正确选“A”，错误选“B”；其余为单项选择题。',bank)

    add_p(doc,'CSPJ 大题标题','一、单项选择题（共 15 题，每题 2 分，共计 30 分；每题有且仅有一个正确选项）',bank,True)
    for n in range(1,16):
        blocks,opts=parse_q1_15(secs[n],n)
        numbered=False
        for kind,val in blocks:
            if kind=='text':
                if not numbered:
                    add_p(doc,'CSPJ 题目',f'{n}. {val}',bank,True)
                    numbered=True
                else:
                    add_p(doc,'CSPJ 题目',val,bank,True)
            elif kind=='code':
                # 若极少数题目一上来就是代码，也先输出题号，避免题号丢失。
                if not numbered:
                    add_p(doc,'CSPJ 题目',f'{n}.',bank,True)
                    numbered=True
                add_code_table(doc,refs,val)
        add_option_table(doc,refs,opts,bank)

    read_title=extract_heading(secs[16],'二、','二、阅读程序')
    add_p(doc,'CSPJ 大题标题',read_title,bank,True)
    current_no=16
    for idx,qno in enumerate((16,17,18),1):
        code,judges,choices,note=parse_reading(secs[qno],qno)
        add_p(doc,'CSPJ 小节标题',f'（{idx}） 阅读下列程序，回答问题。',bank,True)
        add_code_table(doc,refs,code)
        if note:
            add_p(doc,'CSPJ 正文',note,bank)
        add_p(doc,'CSPJ 类型标题','·判断题：',bank,True)
        for j,txt in enumerate(judges,1):
            if global_numbering:
                label=f'{current_no}. '
                current_no+=1
            else:
                label=f'({j}). '
            # J 组原卷判断题直接留（ ），不额外显示 A/B 选项
            qtxt=txt
            if group=='J' and not re.search(r'[（(]\s*[）)]\s*$',qtxt):
                qtxt += '（ ）'
            add_p(doc,'CSPJ 题目',label+qtxt,bank,False if group=='J' else True)
            if group!='J':
                add_option_table(doc,refs,['A. 正确','B. 错误'],bank)
        add_p(doc,'CSPJ 类型标题','·单选题：',bank,True)
        for k,(txt,opts) in enumerate(choices,1):
            if global_numbering:
                label=f'{current_no}. '; current_no+=1
            else:
                label=f'({len(judges)+k}). '
            add_p(doc,'CSPJ 题目',label+txt,bank,True)
            add_option_table(doc,refs,opts,bank)

    fill_title=extract_heading(secs[19],'三、','三、完善程序（单选题，每小题 3 分，共计 30 分）')
    add_p(doc,'CSPJ 大题标题',fill_title,bank,True)
    for idx,qno in enumerate((19,20),1):
        pre,code,groups=parse_fill(secs[qno],qno)
        blocks=clean_prose(pre)
        # 第一段说明跟在（1）/（2）后，符合真题风格
        if blocks:
            add_p(doc,'CSPJ 正文',f'（{idx}）{blocks[0]}',bank,True)
            for b in blocks[1:]: add_p(doc,'CSPJ 正文',b,bank)
        else:
            add_p(doc,'CSPJ 小节标题',f'（{idx}）',bank,True)
        add_code_table(doc,refs,code)
        for sub,(txt,opts) in enumerate(groups,1):
            if global_numbering:
                label=f'{current_no}. '; current_no+=1
            else:
                label=f'({sub}). '
            add_p(doc,'CSPJ 题目',label+txt,bank,True)
            add_option_table(doc,refs,opts,bank)

    doc.core_properties.author=''; doc.core_properties.last_modified_by=''
    out.parent.mkdir(parents=True,exist_ok=True)
    doc.save(out)
    info(f'完成：{out}')
    info(f'结构检查通过：15 道单选 + 3 个阅读程序块 + 2 个完善程序块；编号模式：{"全局" if global_numbering else "块内"}。')


def main():
    ap=argparse.ArgumentParser(description='CSP 初赛 Markdown 一键转 Word（Paper Studio V0.4，判断题紧凑布局）')
    ap.add_argument('input',type=Path)
    ap.add_argument('-o','--output',type=Path,default=None)
    ap.add_argument('-t','--template',type=Path,default=DEFAULT_TEMPLATE)
    ap.add_argument('--numbering',choices=['auto','global','local'],default='auto',help='阅读/完善题号：auto/J组全局/S组块内')
    args=ap.parse_args()
    src=args.input.expanduser().resolve(); out=(args.output.expanduser().resolve() if args.output else src.with_suffix('.docx'))
    template=args.template.expanduser().resolve()
    try: convert(src,out,template,args.numbering)
    except ConvertError as e:
        print('\n[转换失败]\n'+str(e),file=sys.stderr); sys.exit(1)

if __name__=='__main__': main()
