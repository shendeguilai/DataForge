import hashlib
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from converter import (
    add_option_table,
    choose_option_shape,
    clear_document_body,
    collect_math,
    convert,
    normalize_code_lines,
    setup_refs,
)
import service_app
from service_app import safe_stem
from studio_core import analyze_markdown


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / 'CSP_初赛模板.docx'


class CoreCompatibilityTest(unittest.TestCase):
    EXPECTED_HASHES = {
        'converter.py': '712301f1634af8985a49188550db8ee3bc86dbf19065cbffe302fa53f311ee73',
        'studio_core.py': '3c2b28a8539891581b555ab28f7a0f8fb2e1b37dac3fcb3aed47df7d9d5c9896',
        'CSP_初赛模板.docx': '2f40586b53ac3185c3364adea6a1f4268c24415aeed0681d62e3cf23e43e75c2',
    }
    SAMPLE_EXPECTATIONS = {
        '2022_CSP-J.md': {'warnings': 2, 'formulas': 79, 'paragraphs': 73, 'tables': 38, 'global': True},
        '2023_CSP-J.md': {'warnings': 0, 'formulas': 41, 'paragraphs': 68, 'tables': 40, 'global': True},
        '2024_CSP-J.md': {'warnings': 0, 'formulas': 35, 'paragraphs': 70, 'tables': 38, 'global': True},
        '2025_CSP-S.md': {'warnings': 2, 'formulas': 116, 'paragraphs': 71, 'tables': 48, 'global': False},
    }

    def test_v04_core_and_template_are_unchanged(self):
        for name, expected in self.EXPECTED_HASHES.items():
            digest = hashlib.sha256((ROOT / name).read_bytes()).hexdigest()
            self.assertEqual(expected, digest, name)

    def test_all_reference_samples_keep_analysis_and_word_structure(self):
        for name, expected in self.SAMPLE_EXPECTATIONS.items():
            with self.subTest(sample=name):
                source = ROOT / 'samples' / name
                markdown = source.read_text(encoding='utf-8-sig')
                analysis = analyze_markdown(markdown)
                self.assertEqual(20, len(analysis['modules']))
                self.assertEqual([], analysis['summary_errors'])
                self.assertEqual(0, analysis['counts']['error'])
                self.assertEqual(expected['warnings'], analysis['counts']['warning'])
                self.assertEqual(expected['formulas'], len(collect_math(markdown)))

                with tempfile.TemporaryDirectory(prefix='csp_regression_') as temp_dir:
                    output = Path(temp_dir) / f'{source.stem}.docx'
                    convert(source, output, TEMPLATE, 'auto')
                    document = Document(output)
                    self.assertEqual(expected['paragraphs'], len(document.paragraphs))
                    self.assertEqual(expected['tables'], len(document.tables))
                    paragraphs = [paragraph.text for paragraph in document.paragraphs]
                    if expected['global']:
                        self.assertTrue(any(text.startswith('16. ') for text in paragraphs))
                    else:
                        self.assertTrue(any(text.startswith('(1). ') for text in paragraphs))
                    with zipfile.ZipFile(output) as archive:
                        xml = archive.read('word/document.xml')
                        self.assertIn(b'<m:oMath', xml)

    def test_v04_option_layout_and_compact_judgement_rows(self):
        self.assertEqual('1x4', choose_option_shape(['A. 1', 'B. 2', 'C. 3', 'D. 4']))
        self.assertEqual('2x2', choose_option_shape([
            'A. `n % (i-1) == 0`', 'B. `n % i == 0`', 'C. `n / i == 0`', 'D. `n - i == 0`'
        ]))
        self.assertEqual('4x1', choose_option_shape([
            'A. C++ 中构造来源于同一基类的多个派生类并依次调用不同接口',
            'B. 这是另一个长度相近且需要独占一行显示的中文选项内容',
            'C. 这个选项同样不应被强行压缩到狭窄的两列表格单元格中',
            'D. 最后一个选项也保持足够长度以触发单列排版规则',
        ]))

        refs_doc = Document(TEMPLATE)
        refs = setup_refs(refs_doc)
        output_doc = Document(TEMPLATE)
        clear_document_body(output_doc)
        short = add_option_table(output_doc, refs, ['A. 正确', 'B. 错误'], {})
        self.assertEqual(1, len(short.rows))
        long = add_option_table(output_doc, refs, [
            'A. 这是一个超过两列表格安全宽度的非常长判断选项文本内容',
            'B. 这是另一个超过两列表格安全宽度的非常长判断选项文本内容',
        ], {})
        self.assertEqual(2, len(long.rows))

    def test_source_line_numbers_are_removed_without_losing_indentation(self):
        code = '01 #include <iostream>\n02\n03 int main() {\n04     return 0;\n05 }'
        self.assertEqual([
            '#include <iostream>', '', 'int main() {', '    return 0;', '}'
        ], normalize_code_lines(code))

    def test_filename_sanitization_preserves_chinese_and_removes_paths(self):
        self.assertEqual('2025_CSP-S', safe_stem('C:\\试卷\\2025_CSP-S.md'))
        self.assertEqual('CSP试卷', safe_stem('../..'))
        self.assertEqual('试卷_终稿', safe_stem('试卷:终稿.md'))

    def test_failed_export_cleans_temporary_directory(self):
        temp_root = Path(tempfile.gettempdir())
        before = set(temp_root.glob('csp_studio_*'))
        with self.assertRaises(service_app.ConvertError):
            service_app.build_word('# 缺少题目模块', 'invalid', 'auto')
        self.assertEqual(before, set(temp_root.glob('csp_studio_*')))

    def test_export_capacity_returns_busy_without_starting_conversion(self):
        class RejectingSlots:
            @staticmethod
            def acquire(blocking=False):
                return False

        original = service_app.EXPORT_SLOTS
        service_app.EXPORT_SLOTS = RejectingSlots()
        try:
            with self.assertRaises(service_app.ExportBusy):
                service_app.build_word('# sample', 'busy', 'auto')
        finally:
            service_app.EXPORT_SLOTS = original


if __name__ == '__main__':
    unittest.main()
